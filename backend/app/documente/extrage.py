"""
Extragerea textului dintr-un document incarcat. ISC-5.

Documentele NU parasesc masina. Extragerea e complet locala: pypdf pentru PDF,
python-docx pentru DOCX, decodare directa pentru text. Catre model pleaca doar
fragmentul de clauza necesar analizei, niciodata fisierul.

Limitare cunoscuta si declarata: PDF-urile scanate, fara strat de text, nu sunt
citite. Ar necesita OCR. Documentul e respins explicit, cu mesaj clar, in loc sa
fie procesat gol si sa produca o analiza vida care pare valida.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

MIN_CARACTERE_UTILE = 200


class DocumentIlizibil(Exception):
    """Documentul nu contine text extractibil."""


@dataclass
class Clauza:
    index: int
    text: str


def _din_pdf(continut: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(continut))
    return "\n".join(p.extract_text() or "" for p in reader.pages)


def _din_docx(continut: bytes) -> str:
    import docx

    d = docx.Document(io.BytesIO(continut))
    bucati = [p.text for p in d.paragraphs]
    for tabel in d.tables:  # clauzele stau adesea in tabele
        for rand in tabel.rows:
            bucati.extend(c.text for c in rand.cells)
    return "\n".join(bucati)


def extrage_text(nume_fisier: str, continut: bytes) -> str:
    nume = nume_fisier.lower()
    if nume.endswith(".pdf"):
        text = _din_pdf(continut)
    elif nume.endswith(".docx"):
        text = _din_docx(continut)
    elif nume.endswith((".txt", ".md")):
        text = continut.decode("utf-8", errors="replace")
    else:
        raise DocumentIlizibil(
            "Format nesuportat. Accept PDF, DOCX, TXT si MD."
        )

    text = re.sub(r"[ \t]+", " ", text)
    if len(text.strip()) < MIN_CARACTERE_UTILE:
        raise DocumentIlizibil(
            "Nu am putut extrage text din document. Daca este un PDF scanat, "
            "are nevoie de OCR, pe care nu il fac inca. Nu iti dau o analiza "
            "goala care sa para valida."
        )
    return text


# Cate clauze analizam implicit dintr-un document.
#
# Fiecare clauza costa un apel de generare plus doua de verificare. Pe modelul
# local asta inseamna aproximativ 20 de secunde per clauza, deci 12 clauze fac
# o cerere HTTP sincrona de patru minute, pe care orice client o taie prin
# timeout. Sase clauze incap in ~2 minute si raman utile pentru un demo.
#
# Solutia corecta pentru productie NU este un numar mai mic, ci un job asincron
# cu identificator si interogare de stare. Nu e in MVP, si e notat in README.
MAX_CLAUZE_IMPLICIT = 6


def imparte_in_clauze(text: str, *, min_len: int = 120,
                      maxim: int = MAX_CLAUZE_IMPLICIT) -> list[Clauza]:
    """Sparge documentul in unitati analizabile.

    Preferam granitele naturale, articol sau clauza numerotata, si cadem pe
    paragrafe cand nu exista numerotare. Pastram doar bucatile suficient de
    lungi ca sa contina o norma, nu titluri si semnaturi.
    """
    bucati = re.split(r"\n(?=\s*(?:Art\.?\s*\d|Articolul\s+\d|\d+\.\s+[A-ZĂÂÎȘȚ]))",
                      text)
    if len(bucati) < 3:
        bucati = re.split(r"\n\s*\n", text)

    clauze: list[Clauza] = []
    for b in bucati:
        curat = re.sub(r"\s+", " ", b).strip()
        if len(curat) >= min_len:
            clauze.append(Clauza(index=len(clauze) + 1, text=curat[:1500]))
        if len(clauze) >= maxim:
            break
    return clauze
